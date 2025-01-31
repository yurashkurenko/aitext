import re
from docx import Document

def create_docx_from_text(text_file, output_docx):
    # Создаем новый документ
    
    doc = Document()
    chapter_pattern1=r'^(Глава \d+\.?)'
    chapter_pattern2=r'^(подпункт \d+\.?)'
    # Читаем текстовый файл
    with open(text_file, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    # Переменная для хранения текущего параграфа
    current_paragraph = []

    for line in lines:
#        line = line.strip()

        # Проверяем, является ли строка заголовком главы
        if re.match(chapter_pattern1, line):
            # Если есть предыдущий параграф, добавляем его в документ
            if current_paragraph:
                doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []

            # Добавляем заголовок главы
            doc.add_heading(line, level=1)
            current_paragraph.append(line)
        elif re.match(chapter_pattern2, line):
            if current_paragraph:
                doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []
            # Добавляем заголовок подпункта
            doc.add_heading(line, level=4)
            current_paragraph.append(line)
        else:
            # Добавляем строку к текущему параграфу
            if line:
                current_paragraph.append(line)

#    # Добавляем последний параграф, если он есть
    if current_paragraph:
        doc.add_paragraph(' '.join(current_paragraph))

#    # Сохраняем документ
    doc.save(output_docx)

# Пример использования
text_file = 'output.txt'  # Путь к исходному текстовому файлу
output_docx = 'output.docx'  # Путь к выходному файлу .docx

create_docx_from_text(text_file, output_docx)
