import re
from docx import Document

def create_docx_from_text(text_file, output_docx, chapter_pattern=r'^(Глава \d+\.?)'):
    # Создаем новый документ
    doc = Document()

    # Читаем текстовый файл
    with open(text_file, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    # Переменная для хранения текущего параграфа
    current_paragraph = []

    for line in lines:
#        line = line.strip()

        # Проверяем, является ли строка заголовком главы
        if re.match(chapter_pattern, line):
            # Если есть предыдущий параграф, добавляем его в документ
            if current_paragraph:
                doc.add_paragraph(' '.join(current_paragraph))
#                p = doc.add_paragraph()
#                run = p.add_run('Строка 5')
#                run.add_break()
#                run = p.add_run('Строка 6')
#                doc.add_paragraph()
#		run.add_break()
                current_paragraph = []

            # Добавляем заголовок главы
            doc.add_heading(line, level=1)
            current_paragraph.append(line)
#            p = doc.add_paragraph()
#            run = p.add_run('Строка 5')
#            run.add_break()
#            run = p.add_run('Строка 6')
        else:
            # Добавляем строку к текущему параграфу
            if line:
                current_paragraph.append(line)
#                p = doc.add_paragraph()
#                run = p.add_run('Строка 5')
#                run.add_break()
#                run = p.add_run('Строка 6')
    # Добавляем последний параграф, если он есть
    if current_paragraph:
        doc.add_paragraph(' '.join(current_paragraph))

    # Сохраняем документ
    doc.save(output_docx)

# Пример использования
text_file = 'output.txt'  # Путь к исходному текстовому файлу
output_docx = 'output.docx'  # Путь к выходному файлу .docx

create_docx_from_text(text_file, output_docx)
